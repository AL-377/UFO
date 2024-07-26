import json
import os
from typing import List
from docx import Document
from docx.shared import Inches
from docx2pdf import convert

class DocGenerator:
    def __init__(self, doc_type: str, template_version: int = 4):
        self.doc_type = doc_type
        self.template_version = template_version

    def _load_logs(self, png_prefix: str, png_appendix: str, resp_file: str, log_dir: str):
        # only for template version 4 temporarily
        # TODO: use eval file to filter the logs
        imgs = []
        for file in os.listdir(log_dir):
            if file.endswith(png_appendix) and file.startswith(png_prefix):
                file_path = os.path.join(log_dir, file)
                creation_time = os.path.getctime(file_path)
                imgs.append((file_path, creation_time))
        
        # order by create time
        imgs.sort(key=lambda x: x[1])
        
        # return the img files name
        steps_imgs = [file for file, _ in imgs]
        # task description
        task_description = None
        # load the resp_file for the every step
        steps_description = []
        resp_jsons = open(os.path.join(log_dir, resp_file), 'r').readlines()
        for i, step in enumerate(resp_jsons):
            resp_jsons[i] = json.loads(step)
            if not task_description:
                task_description = resp_jsons[i]['Request']
            if "thought" in resp_jsons[i]:
                steps_description.append(resp_jsons[i]['thought'])
        
        return steps_imgs, task_description, steps_description

    def generate_markdown(self, save_path: str, task_description: str, steps_description: List[str], steps_imgs: List[str]):
        markdown_content = f"# Task Description\n\n{task_description}\n\n"
        
        for i, (step_desc, img) in enumerate(zip(steps_description, steps_imgs)):
            markdown_content += f"## Step {i + 1}\n"
            markdown_content += f"{step_desc}\n"
            markdown_content += f"![Related Image]({img})\n\n"
        
        with open(save_path, 'w') as f:
            f.write(markdown_content)

    def generate_docx(self, save_path: str, task_description: str, steps_description: List[str], steps_imgs: List[str]):
        doc = Document()
        doc.add_heading('Task Description', level=1)
        doc.add_paragraph(task_description)
        
        for i, (step_desc, img) in enumerate(zip(steps_description, steps_imgs)):
            doc.add_heading(f'Step {i + 1}', level=2)
            doc.add_paragraph(step_desc)
            doc.add_picture(img, width=Inches(5))
        
        doc.save(save_path)

    def generate_pdf_from_docx(self, docx_path: str, pdf_path: str):
        convert(docx_path, pdf_path)

    def generate(self, save_path: str, log_dir: str, resp_file: str):
        steps_imgs, task_description, steps_description = self._load_logs(
            png_prefix='action_step',
            png_appendix='selected_controls.png',
            resp_file=resp_file,
            log_dir=log_dir
        )

        if self.doc_type == 'markdown':
            self.generate_markdown(save_path, task_description, steps_description, steps_imgs)
        elif self.doc_type == 'docx':
            self.generate_docx(save_path, task_description, steps_description, steps_imgs)
        elif self.doc_type == 'pdf':
            temp_docx_path = save_path.replace('.pdf', '.docx')
            self.generate_docx(temp_docx_path, task_description, steps_description, steps_imgs)
            self.generate_pdf_from_docx(temp_docx_path, save_path)
            os.remove(temp_docx_path)


if __name__ == '__main__':
    doc_gen = DocGenerator(doc_type='markdown', template_version=4)
    doc_gen.generate(save_path='sample.md', log_dir='logs/test_lam_finish', resp_file='response.log')

    doc_gen = DocGenerator(doc_type='pdf', template_version=4)
    doc_gen.generate(save_path='sample.pdf', log_dir='logs/test_lam_finish', resp_file='response.log')

    doc_gen = DocGenerator(doc_type='docx', template_version=4)
    doc_gen.generate(save_path='sample.docx', log_dir='logs/test_lam_finish', resp_file='response.log')


